from django.shortcuts import render
from django.http import HttpResponse
from django.shortcuts import HttpResponse
from django.http import StreamingHttpResponse

import os

from docx.oxml.ns import qn
from docx import Document
from docx.shared import Inches
import time

import threading


import cv2
import numpy as np

from aip import AipOcr
import time
import sys
import importlib

from docx.oxml.ns import qn
from docx import Document
from docx.shared import Inches

from img_db.models import IMG
from login.models import CustomModels
import urllib

from skimage import io
x = []
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
def upload(request):

    if request.method=='POST':

        
        new_img = IMG(
            img=request.FILES.get('img')
        )

        newtime = int(time.time())
        customid = request.session['customid']
        imgname = '%s_%s.png'%(customid,newtime)
        customid = request.session['customid']

        docxname = '%s_%s.docx'%(customid,newtime)
        request.session['imgname'] = imgname
        request.session['docxname'] = docxname 
        new_img.name = imgname
        new_img.save()



        docxname = './media/download/' + docxname

        t = threading.Thread(target = thread_form,args = (imgname,docxname,),name = 'thread_form')
        t.start()
        return render(request,'waiting.html')
def thread_form(imgname,docxname):
    x_y, html_txt = form(imgname,docxname)
    x.append(x_y)
    x.append(html_txt)
def img_form(request):
    html_txt = x.pop()
    imgname = request.session['imgname']
    img = IMG.objects.get(name = imgname)
    return render(request,'div.html',{'html_txt':html_txt,'img':img})


def download(request):
    def file_iterator(file_name, chunk_size=512):
        with open(file_name) as f:
            while True:
                c = f.read(chunk_size)
                if c:
                    yield c
                else:
                    break
    the_file_name = "demo.docx"
    response = StreamingHttpResponse(file_iterator(the_file_name))
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = 'attachment;filename="{0}"'.format(the_file_name)
    return response
def after_img_form(request):
    x_y = x.pop()
    print(x_y)
    for i in x_y:
        content = request.GET['%s.%s'%(i[0],i[1])]
        i.append(content)
    docxname = request.session['docxname']
    path = os.path.join('./media/download/',docxname)

    the_file_name = path
    build_doc(path,x_y)  
    file = open(path,'rb')
    response = StreamingHttpResponse(file)
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = 'attachment;filename="{0}"'.format(the_file_name)
    
    return response




def test1(request):
    return render(request,'test.html')
def test2(request):
    a = request.GET['123']
    print('普通'+a)
    b = a.replace('&#13;&#10;','\n')
    print('html'+b)
    return HttpResponse('puton : %s\nteshu : %s'%(a,b))


def build_doc(name,x_y_v):
    doc = Document(name)
    for i in x_y_v:
        for table in doc.tables:
            cell = table.cell(i[1],i[0])


            cell.text = i[2].replace('\r','')
    doc.save(name)

def dengdai():
    return HttpResponse('请等待')
































#全局变量


#求交点
class Point(object):
    x =0
    y= 0
    # 定义构造方法
    def __init__(self, x=0, y=0):
        self.x = x
        self.y = y

#直线
class Line(object):
    # a=0
    # b=0
    # c=0
    def __init__(self, p1, p2):
        self.p1 = p1
        self.p2 = p2
#文字识别接口
class wordsOCRDetect(object):
    #初始化秘钥信息
    def initIDKey(self,APP_ID,API_KEY,SECRET_KEY):
        """ 你的 APPID AK SK """
        self.client = AipOcr(APP_ID, API_KEY, SECRET_KEY)

    #读取图片，返回文字结果
    def get_image_results(self,img):
        sleeptime = 0.5 #每次识别的等待时间，防QPS5
        # 调用通用文字识别, 图片参数为本地图片
        resultList = []
        results = self.client.basicGeneral(img)
        #print (results)
        time.sleep(sleeptime)   #该死的百度不让并发执行
        for word in results['words_result']:
            resultList.append(word['words'])

        return resultList

#直线表达式 
def GetLinePara(line):
    line.a =line.p1.y - line.p2.y
    line.b = line.p2.x - line.p1.x
    line.c = line.p1.x *line.p2.y - line.p2.x * line.p1.y
#检查交点是否在线段内
def CheckCrossPoint(l1,l2):
    GetLinePara(l1)
    GetLinePara(l2)
    d = l1.a * l2.b - l2.a * l1.b
    if d==0:
        return False
    p=GetCrossPoint(l1,l2)

    #L1比较
    if l1.p1.x>l1.p2.x:
        max1X=l1.p1.x+4
        min1X=l1.p2.x-4
    else:
        max1X=l1.p2.x+4
        min1X=l1.p1.x-4
    if l1.p1.y>l1.p2.y:
        max1Y=l1.p1.y+4
        min1Y=l1.p2.y-4
    else:
        max1Y=l1.p2.y+4
        min1Y=l1.p1.y-4
    #l2xy
    if l2.p1.x>l2.p2.x:
        max2X=l2.p1.x+4
        min2X=l2.p2.x-4
    else:
        max2X=l2.p2.x+4
        min2X=l2.p1.x-4
    if l2.p1.y>l2.p2.y:
        max2Y=l2.p1.y+4
        min2Y=l2.p2.y-4
    else:
        max2Y=l2.p2.y+4
        min2Y=l2.p1.y-4
    if d==0:
        return False
    elif int(p.x)>max1X or int(p.x)<min1X or int(p.y)>max1Y or int(p.y)<min1Y or int(p.x)>max2X or int(p.x)<min2X or int(p.y)>max2Y or int(p.y)<min2Y:
        '''print('x:',int(p.x),'y:',int(p.y))
        print('min1x:',min1X,'max1x:',max1X,'min1y:',min1Y,'max1y:',max1Y)
        print('min2x:',min2X,'max2x:',max2X,'min2y:',min2Y,'max2y:',max2Y)'''
        return False
    else:
        return True
#得到交点
def GetCrossPoint(l1,l2):
    GetLinePara(l1)
    GetLinePara(l2)
    d = l1.a * l2.b - l2.a * l1.b
    p=Point()
    p.x = int((l1.b * l2.c - l2.b * l1.c)*1.0 / d)
    p.y = int((l1.c * l2.a - l2.c * l1.a)*1.0 / d)
    return p
##结束求交点

##切片排序
#X，Y坐标点集合
def Sort(crosslist):
    X={}
    Y={}
    for crosspoint in crosslist:
        if int(crosspoint[0]) not in X:
            X[int(crosspoint[0])]=[crosspoint]
        else:
            X[int(crosspoint[0])].append(crosspoint)
            X[int(crosspoint[0])].sort(key=lambda x:x[1])
        if int(crosspoint[1]) not in Y:
            Y[int(crosspoint[1])]=[crosspoint]
        else:
            Y[int(crosspoint[1])].append(crosspoint)
            Y[int(crosspoint[1])].sort(key=lambda x:x[0])

    return X,Y
#图像色彩取反
def transform(data):
    for i in range(len(data)):
        for j in range(len(data[0])):
            if data[i,j]==255:
                data[i,j]=0
            else :
                data[i,j]=255
    return data
#读取图片
def readImage(imagepath):
    with open(imagepath, 'rb') as fp:
        img = fp.read()
        return img
#word绘制表格
def doc_paint(hengxian,shuxian,wordname):
    document = Document()
    width = len(hengxian[0])
    height = len(shuxian)
    t = document.add_table(rows = height,cols=width, style = 'Table Grid')
    list_heng = []
    list_shu = []
    inverted_hengxian = []#将横线的矩阵转置，方便合并
    for i in range(len(hengxian[0])):
        inverted_hengxian.append([])
    for i in hengxian:
        count = 0
        for j in i:
            inverted_hengxian[count].append(j)
            count += 1
    x = 0
    for i in inverted_hengxian:#可以优化，找出横线中需要合并的单元格，在list_heng中每个元素[x1,y1,x2,y2]
        y = 0
        lock = 0
        for j in i:
            if j == 0 and lock == 0:
                list_heng.append([])
                list_heng[-1].append(x)
                list_heng[-1].append(y-1)#
                lock = 1
                if i[y+1] == 1:
                    list_heng[-1].append(x)
                    list_heng[-1].append(y)
                    lock = 0
            elif j == 0 and lock ==1:
                if i[y+1] == 1:
                    list_heng[-1].append(x)
                    list_heng[-1].append(y)
                    lock = 0
            y += 1

        x += 1
    y = 0
    for i in shuxian:#可以优化，找出竖线中需要合并的单元格，在list_shu中每个元素[x1,y1,x2,y2]
        x = 0
        lock = 0
        for j in i :
            if j == 0 and lock == 0:
                list_shu.append([])
                list_shu[-1].append(x-1)#
                list_shu[-1].append(y)
                lock = 1
                if i[x+1] == 1:
                    list_shu[-1].append(x)
                    list_shu[-1].append(y)
                    lock = 0               
            elif j == 0 and lock ==1:
                if i[x+1] == 1:
                    list_shu[-1].append(x)
                    list_shu[-1].append(y)
                    lock = 0
            x += 1
        y += 1          
    #print(list_heng,list_shu)
    for i in list_heng:
        row0 = t.rows[i[1]]
        row1 = t.rows[i[3]]
        a = row0.cells[i[0]]
        b = row1.cells[i[0]]
        a.merge(b)
    for i in list_shu:
        row0 = t.rows[i[1]]
        a = row0.cells[i[0]]
        b = row0.cells[i[2]]
        a.merge(b)
    
    document.add_page_break()
    document.save(wordname)
    return (list_heng,list_shu)

#显示图片
def showImages(l,img):

    i=0
    for x1i,y1i,x2i,y2i in l[:]:
        i=i+1
        cv2.line(img,(x1i,y1i),(x2i,y2i),(0,0,255),2)
        print ("i=",i,"x1,y1=",x1i,y1i)
        print("i=", i, "x2,y2=", x2i, y2i)
        print ("")
        result = img.copy()
        #cv2.imshow('Result', result)
        #cv2.waitKey(0)
        #cv2.destroyAllWindows()
def form(imgname,wordname):
    #主程序入口位置
    importlib.reload(sys)
    #计算整个识别过程，用了多少时间
    timestart=round(time.clock(),2)
    print (timestart),(" s")

    #传入账户api秘钥信息
    #APP_ID = '10996912'
    #API_KEY = 'mLyRcBTu5Uz3Qxk2Gajlaxyr'
    #SECRET_KEY = 'GKF9qwwS2u3uVCcYMvjrbhhf6ZqK7e9u'
    
    APP_ID = '11044010'
    API_KEY = 'ov1zBYPHHaGUnmQTkI184n7Z'
    SECRET_KEY = 'EGb767Z9ovo7aSj9X4rwzDLqiN1TPooH'
    
    img1 = IMG.objects.get(name = imgname)
    print(img1.img.url)
    imgpath = '.' + img1.img.url

    print(imgpath)


    
    img = cv2.imread(imgpath, 0) #图片路径

    #cv2.imshow('gray',img)
    #图片尺寸
    size = img.shape
    #使用二值化
    ret,edges = cv2.threshold(img,127,255,cv2.THRESH_BINARY)
    edges=transform(edges)
    #检测线段的端点
    lines = cv2.HoughLinesP(edges,1,np.pi/180,30,minLineLength=150,maxLineGap=4)
    #print(lines.type)
    l = lines[:,0,:]
    #在原图中，标注出所有的线段。
    #showImages(l,img)

    crosspoint=Point()
            
    #交点集合
    crosslist=[]
    for i in range(len(l)):
        p1=Point(l[i,0],l[i,1])
        p2=Point(l[i,2],l[i,3])
        line1=Line(p1,p2)
        for x1i,y1i,x2i,y2i in l[i:,:]:
            p3=Point(x1i,y1i)
            p4=Point(x2i,y2i)
            line2=Line(p3,p4)
            if CheckCrossPoint(line1,line2)==True:
                crosspoint=GetCrossPoint(line1,line2)
                crosslist.append([crosspoint.x,crosspoint.y])
    #print(crosslist)
    #print(len(X))
    X,Y=Sort(crosslist)

    #X，Y索引列表，升序排列
    XkeysTemp=X.keys()
    Xkeys=[]
    Xkeys=XkeysTemp
    Xkeys=sorted(Xkeys)
    #print(Xkeys)
    YkeysTemp=Y.keys()
    Ykeys=[]
    Ykeys=YkeysTemp
    Ykeys=sorted(Ykeys)
    #print(Ykeys)

    #集合：左上交点坐标，右边交点X坐标，下面交点Y坐标
    width_and_height=[]
    for xkey,xvalue, in X.items():
        for i in range(len(X[xkey])-1):
            width=X[xkey][i+1][1]
            for j in range(len(Y[X[xkey][i][1]])-1):
                if Y[X[xkey][i][1]][j]==X[xkey][i]:
                    height=Y[X[xkey][i][1]][j+1][0]
                    width_and_height.append([X[xkey][i][0],X[xkey][i][1],width,height])
                else:
                    continue
    #按交点坐标升序排序
    width_and_height.sort(key=lambda x:x[1])
    width_and_height.sort(key=lambda x:x[0])
    print('width-and-height:\n',width_and_height)
    #print(minWidth,minHeight)
    #print(size)

    #横线 竖线矩阵
    hengxian = [([0] * (len(Xkeys)-1)) for i in range(len(Ykeys))]
    #print(hengxian)
    shuxian = [([0] * (len(Xkeys))) for i in range(len(Ykeys)-1)]
    #print(shuxian)
    #print('lenx= ',(len(Xkeys)),'leny= ',(len(Ykeys)))

    #消除轻微倾斜度的线段端点集合，现在误差5像素以内
    ltemp = l.tolist()
    ltemp.sort(key=lambda x:x[1])
    #print(ltemp)
    #处理端点与索引值取相同
    for i in range(len(ltemp)):
        if(ltemp[i][0] not in Xkeys):
            for j in range(len(Xkeys)):
                if(abs(Xkeys[j]-ltemp[i][0])<5):
                    ltemp[i][0] = Xkeys[j]
        if(ltemp[i][2] not in Xkeys):
            for j in range(len(Xkeys)):
                if(abs(Xkeys[j]-ltemp[i][2])<5):
                    ltemp[i][2] = Xkeys[j]
        if(ltemp[i][1] not in Ykeys):
            for j in range(len(Ykeys)):
                if(abs(Ykeys[j]-ltemp[i][1])<5):
                    ltemp[i][1] = Ykeys[j]
        if(ltemp[i][3] not in Ykeys):
            for j in range(len(Ykeys)):
                if(abs(Ykeys[j]-ltemp[i][3])<5):
                    ltemp[i][3] = Ykeys[j]

    #求横线链接状态矩阵
    k=0
    for i in range(len(ltemp)):
        #print(ltemp[i][0],ltemp[i][1],ltemp[i][2],ltemp[i][3])
        #y相等，为横线
        if(ltemp[i][1]==ltemp[i][3]):
            pos1 = Xkeys.index(ltemp[i][0])
            pos2 = Xkeys.index(ltemp[i][2])
            for j in range(len(hengxian)):
                if(pos1<= j <pos2):
                    hengxian[k][j] = 1
            k=k+1 #换行

    print('横线矩阵：\n',hengxian)
    #重新排序用于求竖线
    ltemp.sort(key=lambda x:x[0])
    #print(len(ltemp))
    #求竖线链接状态矩阵
    k=0
    for i in range(len(ltemp)):
        #print(ltemp[i][0],ltemp[i][1],ltemp[i][2],ltemp[i][3])
        #x相等，为竖线
        if(ltemp[i][0]==ltemp[i][2]):
            pos1 = Ykeys.index(ltemp[i][1])
            pos2 = Ykeys.index(ltemp[i][3])
            if(pos1>pos2):
                for j in range(len(shuxian)):
                    if(pos2<= j <pos1):
                        shuxian[j][k] = 1
            else:
                for j in range(len(shuxian)):
                    if(pos1<= j <pos2):
                        shuxian[j][k] = 1
            k=k+1  #换列
    print('竖线矩阵：\n',shuxian)

    #用横线和竖线矩阵绘制word表格
    doc_paint(hengxian,shuxian,wordname)

    width = len(hengxian[0])
    height = len(shuxian)
    l = doc_paint(hengxian,shuxian,wordname)
    list_heng = l[0]
    list_shu = l[1]
    htmlform = Htmlform()
    html = htmlform.buildform(height,width,list_heng,list_shu) 



    
    #定义一个BaiduOCR检测对象
    ocrDemo=wordsOCRDetect()

    #传入账户api秘钥信息
    ocrDemo.initIDKey(APP_ID,API_KEY,SECRET_KEY)
    
    #切片识别后的文字写入word
    #打开word文档
    doc = Document(wordname) 
    #show切片
    x_y = []
    for i in range(len(width_and_height)):
        #print(width_and_height[i][1],width_and_height[i][2],width_and_height[i][0],width_and_height[i][3])
        x = Xkeys.index(width_and_height[i][0])
        y = Ykeys.index(width_and_height[i][1])

        smallimg = img[width_and_height[i][1]:width_and_height[i][2],width_and_height[i][0]:width_and_height[i][3]]
        #print(smallimg)
        #cv2.imshow('qiege',smallimg)
        #cv2.waitKey(0)
        #print(x,y)
        
        #返回图片中的文字字符串列表
        cv2.imwrite(r'filename.jpg', smallimg)

        #图片路径
        imagepath=r'filename.jpg'      

        colimg=readImage(imagepath)
        
        wordlist=ocrDemo.get_image_results(colimg)
        
        #wordlist = [str(i),str(i)]
        #word表格传入文字

        text = ''
        if wordlist:
            text = wordlist[0]
            for word in wordlist[1:]:

                text += '&#13;&#10;' + word
        
        text = html[y][x] + text + '</textarea></td>'
        #list_text = [text]
        html[y][x] = text
        x_y.append([])
        x_y[i].append(x)
        x_y[i].append(y)
                       
    html_txt = ''
    for i in html:
        html_txt += '<tr>'
        for j in i:
            for k in j:
                html_txt += k
        html_txt += '</tr>'
    '''    
    html_txt = '<!DOCTYPE HTML> <html> <style type=\'text/css\'> textarea{width:98%;overflow-x:visible;overflow-y:visible;} </style> <body> <table border=1>' + html_txt
    print('1')

    print('2')
    html_txt += '</table> <input type="submit" value="提交"/> </body> </html>'
    f = open('../templates/test.html','w')
    f.write('%s'%(html_txt))
    f.close() 
            '''
    doc.save(wordname)
    #输出结束时间
    print ("it cost ",round(time.clock(),2)-timestart," s")
    return(x_y,html_txt)




class Htmlform(object):
    def __init__(self):
        pass
    def buildform(self,height,width,list_heng,list_shu):#height为列数，width为行数
        html = []
        html_txt = ''
        for i in range(height):
            html.append([])
            for j in range(width):
  #之后，下面行代码需要加合适ID以标记
                html[i].append('<td><textarea type="text" name="%s.%s" value="">'%(j,i))#(x,y)的形式
        for i in list_heng:
            x = i[0]
            y1 = i[1]
            y2 = i[3]
            rowspan = y2-y1+1
            html[y1][x] = '<td rowspan=%s><textarea type="text" name="%s.%s" value="">'%(rowspan,x,y1)
            for j in range(y2-y1):
                html[y1+j+1][x] = ' '
        for i in list_shu:
            y = i[1]
            x1 = i[0]
            x2 = i[2]
            colspan = x2-x1+1
            html[y][x1] = '<td colspan=%s><textarea type="text" name="%s.%s" value="">'%(colspan,x1,y)
            for j in range(x2-x1):
                html[y][x1+j+1] = ''
        for i in html:
            html_txt += '<tr>'
            for j in i:

                for k in j:    
                    html_txt += k
            html_txt += '</tr>'
        print(html)
        print(html_txt)
        return html
    def buildhtml(self,height,width,list_heng,list_shu):
        html_txt = '<!DOCTYPE HTML> <html> <style type=\'text/css\'> textarea{width:98%;overflow-x:visible;overflow-y:visible;} </style> <body> <table border=1>'
        print('1')
        html_txt += self.buildform(height,width,list_heng,list_shu)
        print('2')
        html_txt += '</table> <textarea type="submit" value="提交"> </body> </html>'
        f = open('test.html','w')
        f.write('%s'%(html_txt))
        f.close()
def url_to_image(url):
	# download the image, convert it to a NumPy array, and then read
	# it into OpenCV format
	resp = urllib.urlopen(url)
	image = np.asarray(bytearray(resp.read()), dtype="uint8")
	image = cv2.imdecode(image, cv2.IMREAD_COLOR)
 
	# return the image
	return image